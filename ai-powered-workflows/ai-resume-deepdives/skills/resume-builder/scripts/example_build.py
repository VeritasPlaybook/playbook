"""
example_build.py

Minimal reference showing the structural pattern of a resume-builder
build script. Educational only. Do not use in production.

The real build script for your setup needs to handle your specific
.docx template's styles and numbering. Ask Claude in a Cowork thread
to write a real one tailored to your template. See Step 9 in SKILL.md.

What this stub demonstrates:

1. Opening a template with python-docx
2. Stripping the body while preserving section properties
3. Inserting cover-letter paragraphs
4. Inserting a hard page break
5. Inserting resume paragraphs
6. Saving to a target path

What this stub deliberately does NOT do, and why your real script must:

- Preserve the template's title-table block and reposition it at the
  top of page 2. This stub strips everything.
- Use the template's actual style names (for example "Heading 1") for
  section headers. This stub uses default styles.
- Match the template's bullet numId (parse word/numbering.xml). This
  stub uses no bullet numbering at all. If you copy this stub and add
  bullets, you must look up your template's numId or bullets will
  render as plain indented text.
- Detect and skip duplicate portfolio links already present in the
  template title table. This stub avoids the issue by stripping
  everything.

Run from the command line with:

    python example_build.py
"""

from docx import Document
from docx.enum.text import WD_BREAK


def build_minimal(
    template_path: str,
    output_path: str,
    cover_paragraphs: list,
    resume_paragraphs: list,
) -> None:
    """Build a two-page .docx from a template plus content lists.

    Parameters
    ----------
    template_path : str
        Absolute path to the .docx template file.
    output_path : str
        Absolute path where the built .docx will be saved.
    cover_paragraphs : list of str
        Each string becomes one paragraph in the cover letter section.
    resume_paragraphs : list of str
        Each string becomes one paragraph in the resume section.

    Notes
    -----
    A real build script needs to additionally preserve the section
    properties (sectPr), use the template's actual style names, match
    the template's bullet numId, and handle the title table block.
    See Step 9 in SKILL.md for the full list of requirements.
    """
    doc = Document(template_path)

    # Clear existing body paragraphs (this also drops the title table;
    # a real script would preserve it).
    for paragraph in list(doc.paragraphs):
        element = paragraph._element
        if element.getparent() is not None:
            element.getparent().remove(element)

    # Insert cover-letter paragraphs.
    for text in cover_paragraphs:
        doc.add_paragraph(text)

    # Insert a hard page break between cover letter and resume.
    page_break_paragraph = doc.add_paragraph()
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    # Insert resume paragraphs.
    for text in resume_paragraphs:
        doc.add_paragraph(text)

    doc.save(output_path)


if __name__ == "__main__":
    # Replace these paths and contents with your own values to test.
    build_minimal(
        template_path="path/to/Application Template.docx",
        output_path="path/to/output/Company - Role.docx",
        cover_paragraphs=[
            "",
            "",
            "",
            "",
            "Dear Hiring Team,",
            "If you are looking for a tailored opener, this is where it goes.",
            "What makes me your ideal candidate?",
            "Reason 1 ...",
            "Reason 2 ...",
            "Reason 3 ...",
            "Reason 4 ...",
            "If you feel I would be an asset to your mission, please reach out.",
            "Sincerely,",
            "Your Name",
        ],
        resume_paragraphs=[
            "Your Name",
            "City, email, LinkedIn, portfolio link",
            "Core Competencies",
            "Item 1  |  Item 2  |  Item 3  |  Item 4",
            "Experience",
            "Most Recent Role at Company, dates",
            "Bullet 1",
            "Bullet 2",
        ],
    )
